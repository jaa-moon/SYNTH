"""Unified text extraction from multiple file formats.

This module provides extractors for PDF, DOCX, plain text, and image files.
Each extractor implements a common interface so that the CLI can route files
to the correct backend based on extension.

For images, the existing :class:`~synth.core.ocr.DocumentScanner` is used.
For PDFs with no embedded text (scanned documents), pages are rendered to
images and passed through OCR as a fallback.
"""

from __future__ import annotations

import logging
import zipfile
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from synth.core.exceptions import NoTextFoundError, SynthError

logger = logging.getLogger(__name__)


# ── Base interface ────────────────────────────────────────────────────────────


class BaseExtractor(ABC):
    """Abstract base for all text extractors."""

    @abstractmethod
    def extract(self, file_path: Path) -> str:
        """Extract text content from *file_path*.

        Returns:
            The extracted text as a single string.

        Raises:
            SynthError: If extraction fails.
            NoTextFoundError: If the file contains no readable text.
        """
        ...

    @property
    @abstractmethod
    def supported_extensions(self) -> set[str]:
        """Set of lowercase file extensions this extractor handles."""
        ...


# ── Plain text / Markdown ────────────────────────────────────────────────────


class TextExtractor(BaseExtractor):
    """Extract text from plain text and markdown files.

    Simply reads the file contents — no processing needed.
    """

    def extract(self, file_path: Path) -> str:
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        text = path.read_text(encoding="utf-8", errors="replace").strip()

        if not text:
            raise NoTextFoundError(path)

        logger.info("TextExtractor: read %d chars from %s", len(text), path.name)
        return text

    @property
    def supported_extensions(self) -> set[str]:
        return {".txt", ".md", ".markdown", ".rst", ".csv"}


# ── PDF extractor (PyMuPDF) ──────────────────────────────────────────────────


class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using PyMuPDF.

    If the PDF contains embedded text (native/digital PDF), it is extracted
    directly. If the text is empty or very short (scanned PDF), pages are
    rendered as images and passed through the OCR pipeline as a fallback.

    Parameters:
        ocr_languages: Language codes for the OCR fallback.
    """

    # Minimum characters to consider the text extraction successful
    # before falling back to OCR.
    MIN_TEXT_LENGTH = 20

    def __init__(self, ocr_languages: list[str] | None = None) -> None:
        self._ocr_languages = ocr_languages or ["en"]

    def extract(self, file_path: Path) -> str:
        import fitz  # PyMuPDF — deferred import

        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        try:
            doc = fitz.open(str(path))
        except Exception as exc:
            raise SynthError(f"Failed to open PDF '{path.name}': {exc}") from exc

        # Attempt direct text extraction
        pages_text: list[str] = []
        for page in doc:
            page_text = page.get_text("text").strip()
            if page_text:
                pages_text.append(page_text)

        doc.close()

        text = "\n\n".join(pages_text).strip()

        if len(text) >= self.MIN_TEXT_LENGTH:
            logger.info(
                "PDFExtractor: extracted %d chars (text-based) from %s",
                len(text),
                path.name,
            )
            return text

        # Fallback: scanned PDF → render pages to images → OCR
        logger.info(
            "PDFExtractor: text too short (%d chars), falling back to OCR for %s",
            len(text),
            path.name,
        )
        return self._ocr_fallback(path)

    def _ocr_fallback(self, path: Path) -> str:
        """Render each page as an image and run OCR."""
        import fitz
        import numpy as np

        from synth.core.ocr import DocumentScanner

        scanner = DocumentScanner(languages=self._ocr_languages)
        doc = fitz.open(str(path))
        all_text: list[str] = []

        for page_num, page in enumerate(doc):
            # Render at 300 DPI for good OCR quality
            pix = page.get_pixmap(dpi=300)
            img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.h, pix.w, pix.n
            )

            # EasyOCR expects BGR or grayscale; pixmap is RGB(A)
            import cv2

            if img_data.shape[2] == 4:
                img_bgr = cv2.cvtColor(img_data, cv2.COLOR_RGBA2BGR)
            else:
                img_bgr = cv2.cvtColor(img_data, cv2.COLOR_RGB2BGR)

            gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)

            # Run OCR directly on the numpy array
            results = scanner._reader.readtext(gray)  # type: ignore[union-attr]
            lines = [text.strip() for _bbox, text, _conf in results if text.strip()]

            if lines:
                all_text.append(f"[Page {page_num + 1}]\n" + "\n".join(lines))

        doc.close()

        if not all_text:
            raise NoTextFoundError(path)

        combined = "\n\n".join(all_text)
        logger.info(
            "PDFExtractor (OCR fallback): extracted %d chars from %s",
            len(combined),
            path.name,
        )
        return combined

    @property
    def supported_extensions(self) -> set[str]:
        return {".pdf"}


# ── DOCX extractor (python-docx) ─────────────────────────────────────────────


class DOCXExtractor(BaseExtractor):
    """Extract text from Microsoft Word (.docx) files.

    Reads all paragraphs and table cells from the document.
    """

    def extract(self, file_path: Path) -> str:
        from docx import Document  # python-docx — deferred import

        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        try:
            doc = Document(str(path))
        except Exception as exc:
            raise SynthError(
                f"Failed to open DOCX '{path.name}': {exc}"
            ) from exc

        parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)

        if not parts:
            raise NoTextFoundError(path)

        combined = "\n".join(parts)
        logger.info(
            "DOCXExtractor: extracted %d chars from %s", len(combined), path.name
        )
        return combined

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}


# ── Apple Pages extractor ────────────────────────────────────────────────────


class PagesExtractor(BaseExtractor):
    """Extract text from Apple Pages (.pages) archives.

    A ``.pages`` file is a ZIP archive.  This extractor attempts two
    strategies:

    1. Look for ``preview.pdf`` inside the archive and extract text via
       :class:`PDFExtractor`.
    2. If no preview is found, look for the ``index.xml`` and extract
       raw text nodes (best-effort, not guaranteed to work on all versions).

    .. note::
        For the most reliable results, users should export their Pages
        documents to PDF or DOCX before scanning.
    """

    def extract(self, file_path: Path) -> str:
        import tempfile

        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        if not zipfile.is_zipfile(str(path)):
            raise SynthError(
                f"'{path.name}' is not a valid Pages file (not a ZIP archive)"
            )

        with zipfile.ZipFile(str(path), "r") as zf:
            names = zf.namelist()

            # Strategy 1: extract preview.pdf
            pdf_candidates = [
                n for n in names if n.lower().endswith(".pdf")
            ]
            if pdf_candidates:
                pdf_name = pdf_candidates[0]
                logger.info(
                    "PagesExtractor: found '%s' inside %s", pdf_name, path.name
                )
                with tempfile.TemporaryDirectory() as tmpdir:
                    extracted_pdf = Path(tmpdir) / "preview.pdf"
                    extracted_pdf.write_bytes(zf.read(pdf_name))
                    pdf_extractor = PDFExtractor()
                    return pdf_extractor.extract(extracted_pdf)

            # Strategy 2: try to parse index.xml for raw text
            xml_candidates = [
                n for n in names if n.lower().endswith(".xml")
            ]
            if xml_candidates:
                import xml.etree.ElementTree as ET

                for xml_name in xml_candidates:
                    try:
                        tree = ET.fromstring(zf.read(xml_name))
                        texts = [
                            elem.text.strip()
                            for elem in tree.iter()
                            if elem.text and elem.text.strip()
                        ]
                        if texts:
                            combined = "\n".join(texts)
                            logger.info(
                                "PagesExtractor (XML): extracted %d chars from %s",
                                len(combined),
                                path.name,
                            )
                            return combined
                    except ET.ParseError:
                        continue

        raise SynthError(
            f"Could not extract text from '{path.name}'. "
            "Please export the document to PDF or DOCX for best results."
        )

    @property
    def supported_extensions(self) -> set[str]:
        return {".pages"}


# ── Image extractor (wraps DocumentScanner) ──────────────────────────────────


class ImageExtractor(BaseExtractor):
    """Extract text from image files using the existing OCR pipeline.

    This is a thin wrapper around :class:`~synth.core.ocr.DocumentScanner`
    to fit the :class:`BaseExtractor` interface.

    Parameters:
        languages: BCP-47 language codes for EasyOCR (default: ``["en"]``).
    """

    def __init__(self, languages: list[str] | None = None) -> None:
        self._languages = languages or ["en"]
        self._scanner: Any = None  # lazy init

    def _get_scanner(self) -> Any:
        """Lazily initialise the scanner to avoid heavy imports at module load."""
        if self._scanner is None:
            from synth.core.ocr import DocumentScanner

            self._scanner = DocumentScanner(languages=self._languages)
        return self._scanner

    def extract(self, file_path: Path) -> str:
        scanner = self._get_scanner()
        return scanner.extract_text(file_path)

    @property
    def supported_extensions(self) -> set[str]:
        return {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}


# ══════════════════════════════════════════════════════════════════════════════
#  Extractor registry
# ══════════════════════════════════════════════════════════════════════════════


# All supported file extensions (union of all extractors)
ALL_SUPPORTED_EXTENSIONS: set[str] = set()

# Mapping: extension → extractor class
_EXTRACTOR_MAP: dict[str, type[BaseExtractor]] = {}

for _cls in (TextExtractor, PDFExtractor, DOCXExtractor, PagesExtractor, ImageExtractor):
    for _ext in _cls.__new__(_cls).supported_extensions:  # type: ignore[arg-type]
        _EXTRACTOR_MAP[_ext] = _cls
        ALL_SUPPORTED_EXTENSIONS.add(_ext)


def get_extractor(
    file_path: Path,
    *,
    ocr_languages: list[str] | None = None,
) -> BaseExtractor:
    """Return the appropriate extractor for *file_path* based on its extension.

    Args:
        file_path: The file to extract text from.
        ocr_languages: Language codes passed to image/PDF OCR extractors.

    Returns:
        An initialised extractor instance.

    Raises:
        SynthError: If the file type is not supported.
    """
    ext = file_path.suffix.lower()

    if ext not in _EXTRACTOR_MAP:
        supported = ", ".join(sorted(ALL_SUPPORTED_EXTENSIONS))
        raise SynthError(
            f"Unsupported file type '{ext}'. Supported: {supported}"
        )

    extractor_cls = _EXTRACTOR_MAP[ext]

    # Pass languages to extractors that need them
    if extractor_cls in (ImageExtractor,):
        return extractor_cls(languages=ocr_languages)  # type: ignore[call-arg]
    if extractor_cls in (PDFExtractor,):
        return extractor_cls(ocr_languages=ocr_languages)  # type: ignore[call-arg]

    return extractor_cls()
